"""DOCX formatting helpers for the industry-analysis board report.

Built from scratch (no sibling-repo dependency). Provides:
    - set_cell_background: table cell shading
    - add_page_number: page-number field for footer
    - add_section_heading: styled heading (H1 / H2 with Arial + blue)
    - add_para: standard paragraph with formatting controls
    - add_flag_box: callout box (orange for methodology, red for issues)
    - make_table: formatted table with header shading + row striping
    - fmt_pct: format decimal as percentage (2dp)

Uses python-docx only (no Node.js, no external tooling). Page setup is
US Letter, 1" margins, Arial 11pt default, centred page-number footer,
right-aligned grey header with "Industry Analysis Report  |  <period>".
"""

from __future__ import annotations

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell


DEFAULT_FONT = "Arial"
DEFAULT_FONT_SIZE = 11

COLOR_H1 = "1F4E79"
COLOR_H2 = "2E75B6"
COLOR_TABLE_HEADER_BG = "1F4E79"
COLOR_TABLE_HEADER_TEXT = "FFFFFF"
COLOR_TABLE_STRIPE_BG = "F2F2F2"
COLOR_FLAG_ORANGE = "F4B183"
COLOR_FLAG_RED = "F4CCCC"
COLOR_INFO_BLUE = "DEEBF7"
COLOR_GREY = "7F7F7F"


def _set_default_font(doc: DocumentType) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = DEFAULT_FONT
    font.size = Pt(DEFAULT_FONT_SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), DEFAULT_FONT)


def _page_setup(doc: DocumentType) -> None:
    for section in doc.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_header_footer(doc: DocumentType, period_label: str) -> None:
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.text = ""
        run = hp.add_run(f"Industry Analysis Report  |  {period_label}")
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLOR_GREY)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        add_page_number(fp)


def new_document(period_label: str) -> DocumentType:
    """Create a styled Document with default font, page setup, and headers."""
    doc = Document()
    _set_default_font(doc)
    _page_setup(doc)
    _add_header_footer(doc, period_label)
    return doc


def set_cell_background(cell: _Cell, color_hex: str) -> None:
    """Shade a table cell with the given hex colour (no leading #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    """Append a Word PAGE field to the given paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    run.font.name = DEFAULT_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COLOR_GREY)


def add_section_heading(doc: DocumentType, text: str, level: int = 1):
    """Styled heading. level=1 -> H1 (16pt deep blue), level=2 -> H2 (13pt medium blue)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = DEFAULT_FONT
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(COLOR_H1)
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(COLOR_H2)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(COLOR_H2)
    return para


def add_para(doc_or_cell, text: str, bold: bool = False, italic: bool = False,
             size: int | None = None, color: str | None = None):
    """Add a paragraph with formatting. Accepts a Document or a table cell."""
    para = doc_or_cell.add_paragraph()
    run = para.add_run(text)
    run.font.name = DEFAULT_FONT
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    return para


def add_flag_box(doc: DocumentType, text: str, color: str = "orange", title: str | None = None):
    """Insert a single-cell shaded table serving as a callout box."""
    color_map = {
        "orange": COLOR_FLAG_ORANGE,
        "red": COLOR_FLAG_RED,
        "blue": COLOR_INFO_BLUE,
    }
    fill = color_map.get(color, COLOR_FLAG_ORANGE)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_background(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.paragraphs[0].text = ""

    if title:
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(title)
        title_run.font.name = DEFAULT_FONT
        title_run.bold = True
        title_run.font.size = Pt(11)

    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if (i == 0 and not title) else cell.add_paragraph()
        if i == 0 and not title:
            para.text = ""
        run = para.add_run(line)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(10)

    _set_cell_borders(cell, fill)
    doc.add_paragraph()
    return table


def _set_cell_borders(cell: _Cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        b = tc_borders.find(qn(f"w:{edge}"))
        if b is None:
            b = OxmlElement(f"w:{edge}")
            tc_borders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), color_hex.lstrip("#"))


def make_table(doc: DocumentType, headers: list[str], rows: list[list[str]],
               col_widths_inches: list[float] | None = None):
    """Build a formatted table with shaded header row and striped rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.font.name = DEFAULT_FONT
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(COLOR_TABLE_HEADER_TEXT)
        set_cell_background(cell, COLOR_TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        tr_cells = table.rows[1 + r_idx].cells
        for c_idx, value in enumerate(row):
            cell = tr_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run("" if value is None else str(value))
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_background(cell, COLOR_TABLE_STRIPE_BG)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_inches:
        for row in table.rows:
            for i, width in enumerate(col_widths_inches):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def fmt_pct(v, digits: int = 2) -> str:
    """Format a decimal value as a percentage string with `digits` decimals."""
    if v is None:
        return ""
    try:
        return f"{float(v) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return str(v)
